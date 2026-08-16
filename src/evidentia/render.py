"""Phase 6 — rendering.

Sections become a document. Markdown is canonical; HTML and DOCX are
projections of the same structured sections, selected by `output_formats` in
config (D-007).

Deterministic sections are rendered here rather than generated. `reporting_period`
is a fact table and `case_index` is a 1,024-row listing — neither is a language
task, and asking a model to produce either would be spending tokens to introduce
risk.
"""

from __future__ import annotations

import csv
from datetime import UTC, datetime
from pathlib import Path
from typing import Protocol

from evidentia.assembler import SectionPacket
from evidentia.config import ReportConfig
from evidentia.generate import GeneratedSection

CASE_TABLE_PREVIEW = 50


# --------------------------------------------------------------------------
# Deterministic section renderers
# --------------------------------------------------------------------------


def render_header_block(packet: SectionPacket, config: ReportConfig) -> str:
    """Title-page facts. Unknown fields are named as unknown, never invented."""
    p = config.product
    rows = [
        ("Product", p.name),
        ("Application number", p.application_number or "not supplied"),
        ("Sponsor", p.sponsor or "not supplied"),
        ("Report type", config.report_type),
        ("Regulatory basis", config.regulatory_basis or "not supplied"),
        ("Reporting period", f"{packet.period_start} to {packet.period_end}"),
        ("Approved indications", p.indication or "not supplied"),
    ]
    total = packet.evidence.get("total_cases", {}).get("value")
    if total is not None:
        rows.append(("Cases in this interval", str(total)))

    lines = ["| Field | Value |", "| --- | --- |"]
    lines += [f"| {k} | {v} |" for k, v in rows]

    if p.unknown_fields:
        lines.append("")
        lines.append(
            "The following were not supplied with the source data and are "
            "therefore not stated in this report: "
            + "; ".join(p.unknown_fields)
            + "."
        )
    return "\n".join(lines)


def render_case_table(packet: SectionPacket, config: ReportConfig) -> str:
    """Case listing.

    The full listing is written to CSV alongside the report; the document
    carries a preview. A 1,024-row markdown table is unreadable and unreviewable,
    and the regulation asks the report to contain *or link to* the listing.
    """
    item = packet.evidence.get("case_index", {})
    columns = item.get("columns", [])
    total = item.get("row_count", 0)

    rows = getattr(render_case_table, "_rows", [])
    if not rows or not columns:
        return f"Case listing of {total} cases is provided as case_index.csv."

    head = "| " + " | ".join(columns) + " |"
    sep = "| " + " | ".join("---" for _ in columns) + " |"
    body = [
        "| " + " | ".join(str(r.get(c, "")) for c in columns) + " |"
        for r in rows[:CASE_TABLE_PREVIEW]
    ]
    note = (
        f"\nShowing the first {min(CASE_TABLE_PREVIEW, total)} of {total} cases. "
        "The complete listing is provided as `case_index.csv`, ordered by report "
        "date. Every aggregate figure in this report traces to these cases."
    )
    return "\n".join([head, sep, *body]) + "\n" + note


DETERMINISTIC_RENDERERS = {
    "header_block": render_header_block,
    "case_table": render_case_table,
}


# --------------------------------------------------------------------------
# Document renderers
# --------------------------------------------------------------------------


class RenderedSection:
    """A section reduced to title plus markdown body, whatever produced it."""

    def __init__(self, section_id: str, title: str, body: str) -> None:
        self.section_id = section_id
        self.title = title
        self.body = body


class Renderer(Protocol):
    extension: str

    def render(
        self, config: ReportConfig, sections: list[RenderedSection], meta: dict
    ) -> bytes | str: ...


class MarkdownRenderer:
    extension = "md"

    def render(
        self, config: ReportConfig, sections: list[RenderedSection], meta: dict
    ) -> str:
        out = [f"# {config.title}", "", f"## {config.product.name}", ""]
        out += [
            f"*{config.report_type} · {meta.get('period', '')}*",
            "",
            "---",
            "",
            "## Contents",
            "",
        ]
        out += [f"{i}. {s.title}" for i, s in enumerate(sections, 1)]
        out += ["", "---", ""]
        for i, s in enumerate(sections, 1):
            out += [f"## {i}. {s.title}", "", s.body, ""]
        out += ["---", "", "## Provenance", "", _provenance_block(meta)]
        return "\n".join(out)


class HtmlRenderer:
    extension = "html"

    def render(
        self, config: ReportConfig, sections: list[RenderedSection], meta: dict
    ) -> str:
        import markdown as md

        body = md.markdown(
            MarkdownRenderer().render(config, sections, meta),
            extensions=["tables"],
        )
        return (
            "<!doctype html><html><head><meta charset='utf-8'>"
            f"<title>{config.title} — {config.product.name}</title>"
            "<style>"
            "body{font-family:Georgia,serif;max-width:52rem;margin:3rem auto;"
            "padding:0 1.5rem;line-height:1.65;color:#1a1a1a}"
            "h1{font-size:1.9rem;margin-bottom:.2rem}"
            "h2{font-size:1.25rem;margin-top:2.4rem;border-bottom:1px solid #ddd;"
            "padding-bottom:.3rem}"
            "table{border-collapse:collapse;width:100%;font-size:.85rem;"
            "font-family:system-ui,sans-serif;margin:1rem 0}"
            "th,td{border:1px solid #ccc;padding:.35rem .55rem;text-align:left}"
            "th{background:#f4f4f4}"
            "code{background:#f4f4f4;padding:.1rem .3rem}"
            "</style></head><body>" + body + "</body></html>"
        )


class DocxRenderer:
    extension = "docx"

    def render(
        self, config: ReportConfig, sections: list[RenderedSection], meta: dict
    ) -> bytes:
        import io

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading(config.title, level=0)
        sub = doc.add_paragraph(config.product.name)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"{config.report_type} · {meta.get('period', '')}")
        doc.add_page_break()

        doc.add_heading("Contents", level=1)
        for i, s in enumerate(sections, 1):
            doc.add_paragraph(f"{i}. {s.title}", style="List Number")
        doc.add_page_break()

        for i, s in enumerate(sections, 1):
            doc.add_heading(f"{i}. {s.title}", level=1)
            _docx_body(doc, s.body)

        doc.add_page_break()
        doc.add_heading("Provenance", level=1)
        for line in _provenance_block(meta).split("\n"):
            if line.strip():
                doc.add_paragraph(line.replace("- ", "").strip())

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


def _docx_body(doc, body: str) -> None:
    """Write a markdown body into a docx, converting pipe tables to real tables."""
    block: list[str] = []

    def flush_table() -> None:
        rows = [r for r in block if r.strip().startswith("|")]
        cells = [
            [c.strip() for c in r.strip().strip("|").split("|")] for r in rows
        ]
        cells = [c for c in cells if not all(set(x) <= set("-: ") for x in c)]
        if not cells:
            return
        table = doc.add_table(rows=0, cols=len(cells[0]))
        table.style = "Table Grid"
        for ri, row in enumerate(cells):
            wr = table.add_row().cells
            for ci, val in enumerate(row[: len(cells[0])]):
                wr[ci].text = val
                if ri == 0:
                    for run in wr[ci].paragraphs[0].runs or [
                        wr[ci].paragraphs[0].add_run("")
                    ]:
                        run.bold = True
        block.clear()

    for line in body.split("\n"):
        if line.strip().startswith("|"):
            block.append(line)
            continue
        if block:
            flush_table()
        if line.strip():
            doc.add_paragraph(line.strip())
    if block:
        flush_table()


def _provenance_block(meta: dict) -> str:
    return "\n".join(
        f"- **{k}**: {v}"
        for k, v in meta.items()
        if k not in {"period"} and v not in (None, "")
    )


RENDERERS: dict[str, Renderer] = {
    "markdown": MarkdownRenderer(),
    "html": HtmlRenderer(),
    "docx": DocxRenderer(),
}


# --------------------------------------------------------------------------
# Assembly
# --------------------------------------------------------------------------


def build_sections(
    config: ReportConfig,
    packets: dict[str, SectionPacket],
    generated: dict[str, GeneratedSection],
) -> list[RenderedSection]:
    """Reduce every configured section to title plus markdown body, in order."""
    out: list[RenderedSection] = []
    for section in config.sections:
        packet = packets[section.id]
        if section.mode == "deterministic":
            fn = DETERMINISTIC_RENDERERS.get(section.renderer or "")
            if fn is None:
                raise ValueError(
                    f"section '{section.id}' has unknown renderer "
                    f"{section.renderer!r}; known: {sorted(DETERMINISTIC_RENDERERS)}"
                )
            body = fn(packet, config)
        else:
            gen = generated.get(section.id)
            if gen is None:
                raise ValueError(f"section '{section.id}' was never generated")
            body = gen.prose
        out.append(RenderedSection(section.id, section.title, body))
    return out


def write_case_index_csv(rows: list[dict], columns: list[str], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=columns, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def render_report(
    config: ReportConfig,
    sections: list[RenderedSection],
    meta: dict,
    out_dir: Path,
    stem: str = "report",
) -> list[Path]:
    """Write every configured output format. Returns the paths written."""
    out_dir.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for fmt in config.output_formats:
        renderer = RENDERERS.get(fmt)
        if renderer is None:
            raise ValueError(
                f"unknown output format {fmt!r}; known: {sorted(RENDERERS)}"
            )
        content = renderer.render(config, sections, meta)
        path = out_dir / f"{stem}.{renderer.extension}"
        if isinstance(content, bytes):
            path.write_bytes(content)
        else:
            path.write_text(content, encoding="utf-8")
        written.append(path)
    return written


def utc_stamp() -> str:
    return datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ")
